"""
FastAPI App: Word to Excel Converter v3.1
Fixed: Support multiple question formats (Brand Health + Customer Experience)
"""

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
import docx
import pandas as pd
import json
import io
import re
from groq import Groq
import traceback

app = FastAPI(title="Word to Excel Converter v3.1")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

MAX_FILE_SIZE = 4.5 * 1024 * 1024  # 4.5MB


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text and tables from .docx file"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        extracted_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text.append(para.text)
        
        for table in doc.tables:
            extracted_text.append("\n=== TABLE ===")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                extracted_text.append(" | ".join(row_data))
            extracted_text.append("=== END TABLE ===\n")
        
        return "\n".join(extracted_text)
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading Word file: {str(e)}")


def detect_structure(text: str) -> dict:
    """
    Detect number of questions, topics and test name from Word
    FIXED: Support multiple question formats
    """
    # COUNT QUESTIONS - Support multiple patterns
    question_patterns = [
        r'^\s*\*?\*?\s*(\d+)\.\s+.+\?',  # Format: **1. Question?**
        r'###\s*\*?\*?\s*Question\s+(\d+):',  # Format: ### **Question 1:**
        r'^\s*Question\s+(\d+)[:：]',          # Format: Question 1: or Question 1：
        r'^\s*Question\s+(\d+)[:：]',     # Format: Question 1:
    ]
    
    all_question_numbers = []
    for pattern in question_patterns:
        matches = re.findall(pattern, text, re.MULTILINE | re.IGNORECASE)
        if matches:
            all_question_numbers.extend([int(m) for m in matches])
    
    # Remove duplicates and count unique questions
    unique_questions = sorted(set(all_question_numbers))
    num_questions = len(unique_questions) if unique_questions else 20
    
    print(f"✓ Detected question numbers: {unique_questions}")
    print(f"✓ Total unique questions: {num_questions}")
    
    # COUNT TOPICS
    topic_patterns = [
        r'Section\s+(\d+)\.',
        r'Part\s+(\d+)\.',
        r'Section\s+(\d+)\.',
    ]
    
    all_topic_numbers = []
    for pattern in topic_patterns:
        matches = re.findall(pattern, text, re.MULTILINE | re.IGNORECASE)
        if matches:
            all_topic_numbers.extend([int(m) for m in matches])
    
    unique_topics = sorted(set(all_topic_numbers))
    num_topics = len(unique_topics) if unique_topics else 6
    
    print(f"✓ Detected topics: {unique_topics}")
    print(f"✓ Total topics: {num_topics}")
    
    # EXTRACT TEST NAME
    test_name_patterns = [
        r'(?:\*\*)?Test\s+Assessment\s+\[(.+?)\]',  # Match [Brand Health]{.mark}
        r'(?:\*\*)?Test\s+Assessment\s+(.+?)(?:\*\*|\(|\[|$)',  # Match "Test Assessment X"
        r'(?:\*\*)?Test\s+Evaluation\s+(.+?)(?:\*\*|\(|$)',
        r'(?:\*\*)?Assessment\s+Test\s+(.+?)(?:\*\*|\(|\n)',
        r'(?:\*\*)?(.+?Test.*?Evaluation.*?)(?:\*\*|\(|\n)',
        r'(?:\*\*)?Test\s+Assessment\s+(.+?)(?:\*\*|\(|$)',
        r'(?:\*\*)?(.+?)\s+(?:Check|Assessment|Health\s+Check)(?:\*\*|\s+\(|\n|$)',
        r'^(?:\*\*)?(.+?Test.+?)(?:\*\*|\n)',
        r'^\*\*(.+?)\*\*',
    ]
    
    detected_test_name = None
    lines = text.split('\n')[:15]  # Check first 15 lines
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 5:
            continue
            
        for pattern in test_name_patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                detected_test_name = match.group(1).strip()
                # Clean up
                detected_test_name = re.sub(r'\*+', '', detected_test_name)
                detected_test_name = re.sub(r'\[.*?\]', '', detected_test_name)
                detected_test_name = re.sub(r'\(.*?\)', '', detected_test_name)
                detected_test_name = re.sub(r'\{.*?\}', '', detected_test_name)  # Add: Remove {.mark}
                detected_test_name = detected_test_name.strip()
                
                # Check reasonable length
                if 5 <= len(detected_test_name) <= 100:
                    print(f"✓ Detected test name: '{detected_test_name}' from line: '{line[:80]}'")
                    break
        
        if detected_test_name:
            break
    
    # Fallback
    if not detected_test_name:
        detected_test_name = 'Business Assessment'
        print(f"⚠ Could not detect test name, using fallback: '{detected_test_name}'")
    
    return {
        'num_questions': num_questions,
        'num_topics': num_topics,
        'detected_test_name': detected_test_name
    }


def analyze_with_llm(text: str, api_key: str, language: str 
                     , structure: dict, metadata: dict) -> list:
    """
    Use LLM to analyze and convert
    FIXED: Improved prompt to handle multiple formats
    """
    try:
        num_questions = structure['num_questions']
        
        # Split into batches if too many questions
        if num_questions > 15:
            print(f"Splitting {num_questions} questions into batches...")
            all_data = []
            batch_size = 10
            
            for i in range(0, num_questions, batch_size):
                batch_questions = min(batch_size, num_questions - i)
                print(f"Processing batch: questions {i+1} to {i+batch_questions}")
                
                batch_structure = {
                    'num_questions': batch_questions,
                    'num_topics': structure['num_topics']
                }
                
                batch_data = _call_llm_api(
                    text, api_key, language,
                    batch_structure, metadata, i+1
                )
                all_data.extend(batch_data)
            
            return all_data[:num_questions]
        else:
            return _call_llm_api(
                text, api_key, language, 
                structure, metadata, 1
            )
    
    except Exception as e:
        error_msg = f"Error calling API: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        raise HTTPException(status_code=500, detail=error_msg)


def _call_llm_api(text: str, api_key: str, language: str, 
                  structure: dict, metadata: dict, start_question: int = 1) -> list:
    """Helper function to call LLM API - FIXED prompt"""
    num_questions = structure['num_questions']
    lang_full = 'Vietnamese' if language == 'vi' else 'English'
    detected_test_name = structure.get('detected_test_name', 'Business Assessment')
    
    # Add: Function to make hashtag
    def make_hashtag(name: str) -> str:
        hashtag = re.sub(r'[^\w\s]', '', name.lower())
        hashtag = re.sub(r'\s+', '_', hashtag)
        return f"#{hashtag}"
    
    # UNIFIED METADATA - All questions use the same
    unified_test_name = metadata.get('test_name') or detected_test_name
    unified_test_category = metadata.get('test_category') or unified_test_name
    
    unified_test_hashtag = metadata.get('test_hashtag')
    if unified_test_hashtag and not unified_test_hashtag.startswith('#'):
        unified_test_hashtag = f"#{unified_test_hashtag}"
    if not unified_test_hashtag:
        unified_test_hashtag = make_hashtag(unified_test_name)
    
    unified_test_cost = metadata.get('test_cost', '0')
    
    # Topic metadata
    unified_topic_name = metadata.get('topic_name') or unified_test_name
    unified_topic_category = metadata.get('topic_category') or unified_test_category
    
    unified_topic_hashtag = metadata.get('topic_hashtag')
    if unified_topic_hashtag and not unified_topic_hashtag.startswith('#'):
        unified_topic_hashtag = f"#{unified_topic_hashtag}"
    if not unified_topic_hashtag:
        unified_topic_hashtag = unified_test_hashtag
    
    # Question metadata
    unified_question_category = metadata.get('question_category') or unified_test_category
    
    unified_question_hashtag = metadata.get('question_hashtag')
    if unified_question_hashtag and not unified_question_hashtag.startswith('#'):
        unified_question_hashtag = f"#{unified_question_hashtag}"
    if not unified_question_hashtag:
        unified_question_hashtag = unified_test_hashtag
    
    # Reference link
    reference_link = metadata.get('reference_link_url', '')
    
    # Default comments
    if language == 'vi':
        default_good = "Well done!"
    else:
        default_good = "Well done!"
    
    # Get relevant text - EXPAND to include text before question
    question_patterns = [
        rf'^\s*\*?\*?\s*{start_question}\.\s+',
        rf'###\s*\*?\*?\s*Question\s+{start_question}:',
        rf'^\s*Question\s+{start_question}[:：]',
    ]
    
    relevant_text = None
    for pattern in question_patterns:
        match = re.search(pattern, text, re.MULTILINE)
        if match:
            # Get from this position to end, EXPAND MORE to fit 5 answers
            start_pos = max(0, match.start() - 800)
            relevant_text = text[start_pos:][:18000]  # Increase to 18000 characters for 5 answers
            break
    
    if not relevant_text:
        relevant_text = text[:18000]
    
    # IMPROVED PROMPT - Support multiple formats
    english_instruction = ""
    if language == 'en':
        english_instruction = """
    * SPECIAL FOR ENGLISH: Extract ONLY the English terms in parentheses from the reading suggestions
    * Example: "Brand Recognition (Brand Recognition), Consistency (Brand Consistency)" → Extract: "Brand Recognition, Brand Consistency"
    * Remove all non-English text, keep only English terms separated by commas
    * If a term has no parentheses (like "Brand Recognition"), keep it as is
    * CRITICAL: PRESERVE ALL line breaks - use \\n for each new line or bullet point
    * Keep bullet points (-) or (•) if they exist"""
    else:
        english_instruction = """
    * FOR VIETNAMESE: Keep the full text including both Vietnamese and English parts
    * CRITICAL: PRESERVE ALL line breaks - use \\n for each new line or bullet point
    * Keep bullet points (-) or (•) if they exist"""
        
    prompt = f"""
Analyze the Word document and extract questions {start_question} to {start_question + num_questions - 1}.

CRITICAL: Return ONLY a pure JSON array [ ... ] with NO markdown, NO text.

Each object must have ALL 31 keys:
{{
  "language": "{language}",
  "test_name": "{unified_test_name}",
  "test_category": "{unified_test_category}",
  "test_hashtag": "{unified_test_hashtag}",
  "test_created_at": "",
  "test_updated_at": "",
  "test_cost": "{unified_test_cost}",
  "topic_name": "{unified_topic_name}",
  "topic_category": "{unified_topic_category}",
  "topic_hashtag": "{unified_topic_hashtag}",
  "topic_created_at": "",
  "topic_updated_at": "",
  "question_text": "...",
  "question_category": "{unified_question_category}",
  "question_hashtag": "{unified_question_hashtag}",
  "question_created_at": "",
  "question_updated_at": "",
  "answer1_text": "...",
  "answer1_info": "improve/review/good",
  "answer2_text": "...",
  "answer2_info": "improve/review/good",
  "answer3_text": "...",
  "answer3_info": "improve/review/good",
  "answer4_text": "...",
  "answer4_info": "improve/review/good",
  "answer5_text": "...",
  "answer5_info": "improve/review/good",
  "good_comment": "{default_good}",
  "improve_comment": "...",
  "read_more_comment": "...",
  "reference_link_url": "{reference_link}"
}}

RULES:
- ALL metadata values are provided above and MUST be used exactly as shown
- CRITICAL: test_name, test_category, test_hashtag MUST BE IDENTICAL for all questions
- Extract COMPLETE question_text in {lang_full}
- Look for answers marked with numbers or letters
- For 5-answer questions: Classify as "improve" (1-worst), "improve" (2), "review" (3), "review" (4), "good" (5-best)
- For 3-answer questions: Classify as "improve" (A-weak), "review" (B-moderate), "good" (C-strong)
- good_comment is ALWAYS: "{default_good}"
- reference_link_url is ALWAYS: "{reference_link}"
- improve_comment and read_more_comment MUST BE IDENTICAL, extracted from reading suggestions:
  * Look for: "Suggestions for further reading:", "**Suggestions for further reading:**", "**Suggestions for further study:**"
  * Extract everything after this marker until the next question starts{english_instruction}
  * CRITICAL LINE BREAK RULES:
    - MUST preserve ALL line breaks from the original document
    - Each bullet point MUST be on a separate line with \\n
    - Each paragraph break MUST use \\n
    - Example: "- Item 1\\n- Item 2\\n- Item 3"
    - Example: "Paragraph 1\\n\\nParagraph 2"
    - If source has numbered lists (1., 2., 3.), preserve with \\n after each
    - If source has bullet points (-, •, *, ◦), preserve with \\n after each
  * IMPORTANT: DO NOT include any emojis, icons, or special symbols (📌, 📖, ✅, 🔹, etc.) in the extracted text
  * Remove all emojis and icons from the extracted content
  * SEARCH CAREFULLY between last answer and the next question
  * If not found, use "" for both fields
- Empty strings for unused answer fields and all *_created_at/*_updated_at fields

Text to analyze:
{relevant_text}
"""
    
    # Calculate max_tokens - increase for 5 answers
    estimated_tokens = num_questions * 800 + 1500  # Increase from 600 to 800 tokens/question
    max_tokens = min(estimated_tokens, 16000)
    
    print(f"Requesting {max_tokens} tokens for {num_questions} questions")
    
    # Call API
    try:       
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "Return ONLY valid, complete JSON array."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=max_tokens
        )
        result_text = response.choices[0].message.content.strip()    

    except Exception as api_error:
        error_msg = str(api_error)
        
        # Handle rate limit errors - CHỈ GIỮ GROQ
        if 'rate_limit' in error_msg.lower() or '429' in error_msg:
            raise ValueError(
                "❌ Groq API quota exceeded (100k tokens/day). "
                "Solutions:\n"
                "1. Wait 34 minutes for quota reset\n"
                "2. OR upgrade Groq to Dev Tier at: https://console.groq.com/settings/billing"
            )
        
        raise ValueError(f"❌ Error calling Groq API: {error_msg}")
    
    # Process JSON
    result_text = result_text.strip()
    
    # Remove markdown
    if result_text.startswith("```json"):
        result_text = result_text[7:].strip()
    if result_text.startswith("```"):
        result_text = result_text[3:].strip()
    if result_text.endswith("```"):
        result_text = result_text[:-3].strip()
    
    result_text = result_text.strip()
    
    # Fix truncated JSON
    if not result_text.endswith(']'):
        print("WARNING: JSON truncated, attempting fix...")
        last_complete = result_text.rfind('},')
        if last_complete > 0:
            result_text = result_text[:last_complete+1] + ']'
        else:
            raise ValueError("JSON truncated and cannot be fixed. Try reducing number of questions.")
    
    # Parse JSON
    try:
        data = json.loads(result_text)
    except json.JSONDecodeError as je:
        raise ValueError(f"Invalid JSON: {str(je)}")
    
    if isinstance(data, dict):
        data = [data]
    
    if not isinstance(data, list) or not data:
        raise ValueError("AI returned invalid data")
    
    # Add this section - Ensure consistent metadata
    unified_test_name = metadata.get('test_name') or structure.get('detected_test_name', 'Business Assessment')
    unified_test_category = metadata.get('test_category') or unified_test_name
    
    unified_test_hashtag = metadata.get('test_hashtag')
    if unified_test_hashtag and not unified_test_hashtag.startswith('#'):
        unified_test_hashtag = f"#{unified_test_hashtag}"
    if not unified_test_hashtag:
        unified_test_hashtag = make_hashtag(unified_test_name)
    
    unified_test_cost = metadata.get('test_cost', '0')
    unified_topic_name = metadata.get('topic_name') or unified_test_name
    unified_topic_category = metadata.get('topic_category') or unified_test_category
    
    unified_topic_hashtag = metadata.get('topic_hashtag')
    if unified_topic_hashtag and not unified_topic_hashtag.startswith('#'):
        unified_topic_hashtag = f"#{unified_topic_hashtag}"
    if not unified_topic_hashtag:
        unified_topic_hashtag = unified_test_hashtag
    
    # Apply to ALL items
    for item in data:
        item['test_name'] = unified_test_name
        item['test_category'] = unified_test_category
        item['test_hashtag'] = unified_test_hashtag
        item['test_cost'] = unified_test_cost
        item['topic_name'] = unified_topic_name
        item['topic_category'] = unified_topic_category
        item['topic_hashtag'] = unified_topic_hashtag
    
    # Post-process: Ensure line breaks and remove icons/emojis
    for item in data:
        # Function to remove emojis and icons
        def remove_emojis(text):
            if not text:
                return text
            import re
            # Pattern to remove Unicode emojis and icons
            emoji_pattern = re.compile(
                "["
                u"\U0001F600-\U0001F64F"  # emoticons
                u"\U0001F300-\U0001F5FF"  # symbols & pictographs
                u"\U0001F680-\U0001F6FF"  # transport & map symbols
                u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
                u"\U00002702-\U000027B0"  # dingbats
                u"\U000024C2-\U0001F251"
                u"\U0001F900-\U0001F9FF"  # Supplemental Symbols and Pictographs
                u"\U0001FA00-\U0001FA6F"  # Chess Symbols
                u"\U00002600-\U000026FF"  # Miscellaneous Symbols
                u"\U00002700-\U000027BF"  # Dingbats
                "]+", flags=re.UNICODE
            )
            text = emoji_pattern.sub('', text)
            
            # Remove special characters
            text = re.sub(r'[📌📖✅📹💡⚠️✔❌🎯🚀📊📈📉💰🎪🌭🔥💪🏆]', '', text)
            
            # Do NOT remove extra spaces here to preserve line breaks
            return text.strip()
        
        # Apply to all text fields
        for key in item.keys():
            if isinstance(item[key], str) and item[key]:
                # Remove emojis
                item[key] = remove_emojis(item[key])
                
                # CRITICAL: Handle line breaks for improve_comment and read_more_comment
                if key in ['improve_comment', 'read_more_comment']:
                    # Convert all line break forms
                    item[key] = item[key].replace('\\n', '\n')  # \\n -> \n
                    item[key] = item[key].replace('\\r\\n', '\n')  # Windows line break
                    item[key] = item[key].replace('\\r', '\n')  # Old Mac line break
                    
                    # Remove extra spaces but KEEP line breaks
                    lines = item[key].split('\n')
                    lines = [line.strip() for line in lines]
                    item[key] = '\n'.join(lines)
    
    print(f"✓ Successfully parsed {len(data)} questions")
    return data


def convert_json_to_excel(data: list) -> bytes:
    """Convert JSON to Excel - FIXED: Ensure consistent metadata"""
    try:
        if not data or not isinstance(data, list):
            raise ValueError("Invalid data")
        
        df = pd.DataFrame(data)
        
        # 31 standard columns
        required_columns = [
            "language", "test_name", "test_category", "test_hashtag",
            "test_created_at", "test_updated_at", "test_cost",
            "topic_name", "topic_category", "topic_hashtag",
            "topic_created_at", "topic_updated_at",
            "question_text", "question_category", "question_hashtag",
            "question_created_at", "question_updated_at",
            "answer1_text", "answer1_info", "answer2_text", "answer2_info",
            "answer3_text", "answer3_info", "answer4_text", "answer4_info",
            "answer5_text", "answer5_info",
            "good_comment", "improve_comment", "read_more_comment",
            "reference_link_url"
        ]
        
        # Add missing columns
        for col in required_columns:
            if col not in df.columns:
                df[col] = ""
        
        # CRITICAL FIX: Ensure test_name, test_category, test_hashtag are the same
        if len(df) > 0:
            # Get values from first row
            first_test_name = df.loc[0, 'test_name']
            first_test_category = df.loc[0, 'test_category']
            first_test_hashtag = df.loc[0, 'test_hashtag']
            first_test_cost = df.loc[0, 'test_cost']
            first_topic_name = df.loc[0, 'topic_name']
            first_topic_category = df.loc[0, 'topic_category']
            first_topic_hashtag = df.loc[0, 'topic_hashtag']
            
            # Apply to ALL rows
            df['test_name'] = first_test_name
            df['test_category'] = first_test_category
            df['test_hashtag'] = first_test_hashtag
            df['test_cost'] = first_test_cost
            df['topic_name'] = first_topic_name
            df['topic_category'] = first_topic_category
            df['topic_hashtag'] = first_topic_hashtag
        
        # Reorder columns
        df = df[required_columns]
        
        # Convert test_cost to number
        def convert_to_number(value):
            if pd.isna(value) or value == '' or value is None:
                return 0
            try:
                cleaned = str(value).replace(',', '').replace(' ', '').strip()
                return float(cleaned) if cleaned else 0
            except:
                return 0
        
        df['test_cost'] = df['test_cost'].apply(convert_to_number)
        
        # Create Excel
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Questions')
            
            worksheet = writer.sheets['Questions']
            from openpyxl.styles import Alignment
            
            for idx, col in enumerate(required_columns):
                if len(df) > 0:
                    max_length = max(
                        df[col].astype(str).apply(len).max(),
                        len(col)
                    )
                else:
                    max_length = len(col)
                
                col_letter = ''
                num = idx + 1
                while num > 0:
                    num -= 1
                    col_letter = chr(65 + (num % 26)) + col_letter
                    num //= 26
                worksheet.column_dimensions[col_letter].width = min(max_length + 2, 50)
                
                if col in ['improve_comment', 'read_more_comment']:
                    worksheet.column_dimensions[col_letter].width = 60
                    for row in range(2, len(df) + 2):
                        cell = worksheet[f'{col_letter}{row}']
                        cell.alignment = Alignment(
                            wrap_text=True, 
                            vertical='top',
                            horizontal='left'
                        )
        
        output.seek(0)
        return output.getvalue()
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating Excel: {str(e)}")


@app.get("/", response_class=HTMLResponse)
async def home():
    """Serve HTML interface"""
    # Inline HTML for Vercel deployment
    html_content = """<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Word → Excel Converter v3 - Custom Metadata</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        .container {
            max-width: 900px;
            margin: 0 auto;
            background: white;
            padding: 40px;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0, 0, 0, 0.3);
        }
        h1 {
            color: #333;
            margin-bottom: 10px;
            font-size: 32px;
            text-align: center;
        }
        .version-badge {
            background: linear-gradient(135deg, #667eea, #764ba2);
            color: white;
            padding: 6px 14px;
            border-radius: 15px;
            font-size: 13px;
            font-weight: 600;
        }
        .subtitle {
            color: #666;
            text-align: center;
            margin-bottom: 40px;
            font-size: 15px;
        }
        button {
            width: 100%;
            padding: 16px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 10px;
            font-size: 17px;
            font-weight: 700;
            cursor: pointer;
            margin-top: 10px;
        }
        button:hover { transform: translateY(-3px); }
        input, select {
            width: 100%;
            padding: 12px;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            margin: 8px 0;
        }
        .form-group { margin-bottom: 20px; }
        label { display: block; margin-bottom: 8px; font-weight: 600; }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 Word → Excel Converter <span class="version-badge">v3.1</span></h1>
        <p class="subtitle">Chuyển đổi file Word sang Excel với Metadata tùy chỉnh</p>
        
        <form id="uploadForm">
            <div class="form-group">
                <label>📄 Chọn file Word (.docx)</label>
                <input type="file" id="file" accept=".docx" required>
            </div>
            
            <div class="form-group">
                <label>🌐 Ngôn ngữ output</label>
                <select id="language">
                    <option value="vi">🇻🇳 Tiếng Việt</option>
                    <option value="en">🇬🇧 English</option>
                </select>
            </div>
            
            <div class="form-group">
                <label>🔑 Groq API Key</label>
                <input type="password" id="apiKey" placeholder="Nhập API key..." required>
                <small>Tạo free tại: <a href="https://console.groq.com/keys" target="_blank">console.groq.com/keys</a></small>
            </div>
            
            <button type="submit">🚀 Chuyển đổi ngay</button>
        </form>
        
        <div id="status" style="margin-top: 20px; padding: 15px; border-radius: 8px; display: none;"></div>
    </div>
    
    <script>
        document.getElementById('uploadForm').addEventListener('submit', async (e) => {
            e.preventDefault();
            const status = document.getElementById('status');
            const formData = new FormData();
            
            formData.append('file', document.getElementById('file').files[0]);
            formData.append('api_key', document.getElementById('apiKey').value);
            formData.append('language', document.getElementById('language').value);
            
            status.style.display = 'block';
            status.style.background = '#e3f2fd';
            status.innerHTML = '⏳ Đang xử lý... Vui lòng đợi (10-30 giây)';
            
            try {
                const response = await fetch('/convert', {
                    method: 'POST',
                    body: formData
                });
                
                if (!response.ok) {
                    const error = await response.json();
                    throw new Error(error.detail || 'Lỗi không xác định');
                }
                
                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = 'converted.xlsx';
                a.click();
                
                status.style.background = '#e8f5e9';
                status.innerHTML = '✅ Chuyển đổi thành công! File đã tải xuống.';
            } catch (error) {
                status.style.background = '#ffebee';
                status.innerHTML = '❌ ' + error.message;
            }
        });
    </script>
</body>
</html>"""
    return HTMLResponse(content=html_content)


@app.post("/convert")
async def convert_word_to_excel(
    file: UploadFile = File(...),
    api_key: str = Form(...),
    language: str = Form('vi'),
    test_name: str = Form(''),
    test_category: str = Form(''),
    test_hashtag: str = Form(''),
    test_cost: str = Form('0'),
    topic_name: str = Form(''),
    topic_category: str = Form(''),
    topic_hashtag: str = Form(''),
    question_category: str = Form(''),
    question_hashtag: str = Form(''),
    reference_link_url: str = Form('')
):
    """Convert Word to Excel with custom metadata"""
    
    # Validate
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files accepted")
    
    file_content = await file.read()
    
    if len(file_content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large (max 4.5MB)")
    
    if not api_key or len(api_key) < 20:
        raise HTTPException(status_code=400, detail="Invalid API key")
    
    # Prepare metadata
    metadata = {
        'test_name': test_name.strip(),
        'test_category': test_category.strip(),
        'test_hashtag': test_hashtag.strip(),
        'test_cost': test_cost.strip(),
        'topic_name': topic_name.strip(),
        'topic_category': topic_category.strip(),
        'topic_hashtag': topic_hashtag.strip(),
        'question_category': question_category.strip(),
        'question_hashtag': question_hashtag.strip(),
        'reference_link_url': reference_link_url.strip()
    }
    
    try:
        # Step 1: Extract text
        print("=" * 80)
        print(f"Processing file: {file.filename}")
        print("=" * 80)
        text = extract_text_from_docx(file_content)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="Word file has no content")
        
        # Debug
        print("\nFIRST 800 CHARS FROM WORD:")
        print(text[:800])
        print("=" * 80)
        
        # Step 2: Detect structure
        print("\nDetecting structure...")
        structure = detect_structure(text)
        print(f"✓ Structure: {structure['num_questions']} questions, {structure['num_topics']} topics")
        print(f"✓ Test name: '{structure.get('detected_test_name')}'")
        
        # Step 3: Analyze with LLM
        structured_data = analyze_with_llm(
            text, api_key, language, structure, metadata
        )
        
        print(f"✓ Got {len(structured_data)} questions from LLM")
        
        # Step 4: Create Excel
        print("\nCreating Excel...")
        excel_content = convert_json_to_excel(structured_data)
        
        # Output filename
        output_filename = file.filename.replace('.docx', f'_converted_{language}.xlsx')
        
        print(f"✓ Success! Output: {output_filename}")
        print("=" * 80)
        
        return StreamingResponse(
            io.BytesIO(excel_content),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename={output_filename}"
            }
        )
    
    except HTTPException:
        raise
    except Exception as e:
        error_detail = f"Error: {str(e)}\n{traceback.format_exc()}"
        print(error_detail)
        raise HTTPException(status_code=500, detail=error_detail)


@app.get("/health")
async def health_check():
    """Health check"""
    return {
        "status": "healthy",
        "version": "3.1",
        "features": ["groq_only", "custom_metadata", "auto_detect"]
    }

app = app